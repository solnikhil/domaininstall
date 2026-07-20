#!/usr/bin/env python3
"""Build the verified domaininstall execution checklist DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "artifacts" / "domaininstall-project-checklist.docx"

# compact_reference_guide token map
FONT = "Calibri"
BODY_SIZE = 11
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
GREEN_FILL = "E2F0D9"
GREEN_TEXT = "2E6B34"
GOLD_FILL = "FFF2CC"
GOLD_TEXT = "7A5A00"
RED_FILL = "FCE4D6"
RED_TEXT = "9B1C1C"


def set_run(run, *, size=BODY_SIZE, color="1F2937", bold=False, italic=False, font=FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != 9360:
        raise ValueError("Table widths must total 9360 DXA")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_table_borders(table, color="D0D5DD", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def set_cell_text(cell, text, *, bold=False, color="1F2937", size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(text), size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run(run, size=8.5, color=MUTED)


def add_bottom_rule(paragraph, color=BLUE, size="14"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_numbering_definition(doc, *, kind):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abs or [0]) + 1
    num_id = max(existing_num or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    set_run(p.add_run(text), size=10.5)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(text), size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return p


def add_status_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [1296, 3312, 3312, 1440])  # 0.9, 2.3, 2.3, 1.0 in
    set_table_borders(table)
    headers = ["STATUS", "ACTION", "EXIT CRITERION", "OWNER / TARGET"]
    for i, label in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], LIGHT_BLUE)
        set_cell_text(table.rows[0].cells[i], label, bold=True, color=INK, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        table.rows[0].cells[i].paragraphs[0].paragraph_format.keep_with_next = True
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])

    for status, action, exit_criterion, owner in rows:
        row = table.add_row()
        set_row_cant_split(row)
        cells = row.cells
        status_fill, status_color = {
            "COMPLETE": (GREEN_FILL, GREEN_TEXT),
            "NEXT": (LIGHT_BLUE, DARK_BLUE),
            "PENDING": (GOLD_FILL, GOLD_TEXT),
            "BLOCKED": (RED_FILL, RED_TEXT),
        }[status]
        set_cell_shading(cells[0], status_fill)
        set_cell_text(cells[0], status, bold=True, color=status_color, size=8.2, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], action, size=9.3)
        set_cell_text(cells[2], exit_criterion, size=9.3)
        set_cell_text(cells[3], owner, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(table, [1296, 3312, 3312, 1440])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_decision_table(doc):
    rows = [
        ("npm package", "domaininstall (not published yet)"),
        ("Primary command", "di"),
        ("Aliases", "domaininstall, dnstall"),
        ("Live mapping", "zuraai.xyz → npm:zuraai"),
        ("DNS record", '_dnstall.zuraai.xyz TXT "dnstall=pkg:npm/zuraai"'),
        ("Current branch", "feat/v0 @ aef8a97 (pushed, clean)"),
        ("GitHub", "Private; default branch is feat/v0; main is behind"),
        ("Planned release", "0.0.1"),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2700, 6660])
    set_table_borders(table, color="D8DEE8")
    for i, label in enumerate(("DECISION", "CURRENT CHOICE / STATE")):
        set_cell_shading(table.rows[0].cells[i], LIGHT_BLUE)
        set_cell_text(table.rows[0].cells[i], label, bold=True, color=INK, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        table.rows[0].cells[i].paragraphs[0].paragraph_format.keep_with_next = True
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for label, value in rows:
        row = table.add_row()
        set_row_cant_split(row)
        cells = row.cells
        set_cell_shading(cells[0], LIGHT_GRAY)
        set_cell_text(cells[0], label.upper(), bold=True, color=DARK_BLUE, size=8.7)
        set_cell_text(cells[1], value, size=9.5)
    set_table_geometry(table, [2700, 6660])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_callout(doc, label, text, *, fill=CALLOUT, accent=DARK_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Pt(7)
    p.paragraph_format.right_indent = Pt(7)
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    for edge, color, size in (
        ("left", accent, "18"),
        ("top", "D8DEE8", "4"),
        ("bottom", "D8DEE8", "4"),
        ("right", "D8DEE8", "4"),
    ):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "5")
        border.set(qn("w:color"), color)
        p_bdr.append(border)
    p_pr.append(p_bdr)
    set_run(p.add_run(f"{label}  "), size=10, color=accent, bold=True)
    set_run(p.add_run(text), size=10, color="344054")


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(BODY_SIZE)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color, before, after in (
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 14, 7),
        (3, 12, DARK_BLUE, 10, 5),
    ):
        style = styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    # Quiet running header/footer.
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    set_run(hp.add_run("DOMAININSTALL  •  EXECUTION CHECKLIST"), size=8.5, color=MUTED, bold=True)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    set_run(fp.add_run("Page "), size=8.5, color=MUTED)
    add_field(fp, "PAGE")
    set_run(fp.add_run(" of "), size=8.5, color=MUTED)
    add_field(fp, "NUMPAGES")

    # memo_masthead opening block.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run("DOMAININSTALL"), size=10, color=BLUE, bold=True)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_run(title.add_run("Project Execution Checklist"), size=23, color="111827", bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_run(subtitle.add_run("From proven v0 prototype to a trustworthy public npm release"), size=12.5, color=MUTED)
    for label, value in (
        ("Snapshot", "20 July 2026"),
        ("Current milestone", "Release readiness and npm distribution"),
        ("Source of truth", "github.com/solnikhil/domaininstall • feat/v0"),
    ):
        mp = doc.add_paragraph()
        mp.paragraph_format.space_after = Pt(2)
        set_run(mp.add_run(f"{label}: "), size=9.5, color="344054", bold=True)
        set_run(mp.add_run(value), size=9.5, color="344054")
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(12)
    add_bottom_rule(rule)

    add_callout(
        doc,
        "CURRENT STATUS",
        "The product works end to end: public DNS resolves zuraai.xyz to npm:zuraai, di installs it, and the isolated TOFU pin is verified. The main remaining gap is distribution and release hygiene—the domaininstall package is still unpublished.",
    )

    add_heading(doc, "1. Current decisions", 1)
    add_decision_table(doc)

    completed_heading = add_heading(doc, "2. Completed proof", 1)
    completed_heading.paragraph_format.page_break_before = True
    add_status_table(
        doc,
        [
            ("COMPLETE", "Core zero-runtime-dependency TypeScript CLI", "Build succeeds on a clean install; package-manager handoff uses argv with shell disabled.", "Verified\n20 Jul"),
            ("COMPLETE", "Record format and input validation", "purl and legacy TXT formats parse; scoped packages, versions, domains, and flag-smuggling checks pass.", "19/19\nsmoke"),
            ("COMPLETE", "Live DNS mapping", "Cloudflare and Google DoH return _dnstall.zuraai.xyz → dnstall=pkg:npm/zuraai.", "TTL 300\nDNSSEC —"),
            ("COMPLETE", "True CLI-level end-to-end test", "npm run test:e2e resolves DNS, runs di, installs zuraai, and verifies isolated TOFU state.", "Passing\n20 Jul"),
            ("COMPLETE", "Primary command and package identity", "npm package remains domaininstall; di is primary, with domaininstall and dnstall aliases.", "Decided"),
            ("COMPLETE", "Demo proof", "36.05-second H.264/AAC vertical demo uses the real di zuraai.xyz flow and truthful DNSSEC state.", "Rendered\n4.93 MB"),
            ("COMPLETE", "Remote synchronization", "feat/v0 is clean and pushed at aef8a97.", "Pushed"),
        ],
    )

    add_heading(doc, "3. Release-critical checklist", 1)
    add_callout(doc, "NEXT MILESTONE", "Publish domaininstall@0.0.1, install it from the public registry, and create a reproducible GitHub release.", fill=LIGHT_BLUE)
    add_status_table(
        doc,
        [
            ("NEXT", "Reconcile release branches", "Choose the canonical branch; bring main to the complete feat/v0 state and set the intended GitHub default branch.", "Owner: —\nTarget: —"),
            ("NEXT", "Add README.md", "README includes install, di usage, TXT setup, trust model, limitations, examples, and troubleshooting.", "Owner: —\nTarget: —"),
            ("NEXT", "Add an MIT LICENSE file", "Repository and npm tarball contain the full license text; GitHub detects the license.", "Owner: —\nTarget: —"),
            ("NEXT", "Add GitHub CI", "Every PR runs npm ci, npm test, build, and package dry-run; live E2E runs separately on schedule/manual dispatch.", "Owner: —\nTarget: —"),
            ("NEXT", "Authenticate and publish domaininstall", "npm view domaininstall returns 0.0.1 and npm install -g domaininstall succeeds from a clean environment.", "Owner: Nikhil\nTarget: —"),
            ("PENDING", "Verify all three binaries from npm", "di --version, domaininstall --version, and dnstall --version each execute the published CLI.", "After publish"),
            ("PENDING", "Tag and create GitHub release", "v0.0.1 tag points to the tested release commit; notes include demo, DNS example, and known limitations.", "After publish"),
            ("PENDING", "Prepare repository for discovery", "Decide public/private status; add description, topics, homepage/demo link, and security reporting path.", "Decision\nrequired"),
        ],
    )

    add_heading(doc, "4. Security gates before broad promotion", 1)
    add_callout(
        doc,
        "POSITIONING GUARDRAIL",
        "TOFU protects returning users from a changed domain→package mapping. It does not protect a first-time user from an already-compromised, expired, or mistyped domain. Avoid claiming complete typosquatting prevention.",
        fill=RED_FILL,
        accent=RED_TEXT,
    )
    add_status_table(
        doc,
        [
            ("NEXT", "Escape untrusted terminal output", "Raw TXT values and metadata cannot emit control sequences in verify or install output.", "P0"),
            ("NEXT", "Make resolver failures explicit", "Network failure, timeout, SERVFAIL, NXDOMAIN, and no-record states produce different messages; Google fallback handles transient DNS failures.", "P0"),
            ("NEXT", "Harden the TOFU store", "Corrupt pins fail loudly; writes are atomic and concurrency-safe; recovery instructions are documented.", "P0"),
            ("NEXT", "Reject ambiguous mappings", "Multiple valid npm records produce a clear ambiguity error instead of silently selecting the first value.", "P0"),
            ("PENDING", "Pin immutable artifact evidence", "Store exact resolved version and npm integrity; surface publisher/provenance signals where available.", "P1"),
            ("PENDING", "Add domain-continuity signals", "Evaluate DNSSEC regression, RDAP creation/expiry, nameserver changes, and pin-age policy.", "P1"),
        ],
    )

    add_heading(doc, "5. Publisher and user experience", 1)
    add_status_table(
        doc,
        [
            ("NEXT", "Create a publisher setup guide", "A domain owner can copy one TXT record, verify it with di verify, and understand TTL/version choices in under five minutes.", "Docs"),
            ("NEXT", "Create a 60-second consumer quickstart", "Install → di zuraai.xyz → inspect summary → confirm → understand the pin, with no unexplained jargon.", "Docs"),
            ("PENDING", "Improve package-manager detection", "Walk parent directories and honor packageManager before falling back to current-directory lockfiles/npm.", "P1"),
            ("PENDING", "Tighten CLI argument behavior", "Unknown flags and extra positional arguments fail clearly; version output reads from package metadata.", "P1"),
            ("PENDING", "Clean legacy terminology", "dpm/_dpm language is removed or explicitly marked historical across research and security docs.", "Docs"),
        ],
    )

    add_heading(doc, "6. Launch and validation", 1)
    add_status_table(
        doc,
        [
            ("NEXT", "Choose the launch wedge", "One sentence consistently positions domain-vouched installs for AI agents/slopsquatting, rather than promising universal package security.", "Founder"),
            ("NEXT", "Recruit three real publisher domains", "At least three maintainers publish records for packages they own and complete di verify successfully.", "Validation"),
            ("PENDING", "Create a minimal landing page", "Explains the problem, DNS proof, exact command, limitations, and publisher setup; links to npm and GitHub.", "After publish"),
            ("PENDING", "Instrument launch metrics", "Track registry installs, successful verifications, active mapped domains, repeat installs, pin-change warnings, and setup completion.", "Post-launch"),
            ("PENDING", "Collect trust objections", "Document why users do or do not prefer domain identity over npm scopes/provenance and revise positioning from evidence.", "Research"),
        ],
    )

    add_heading(doc, "7. Release-ready definition", 1)
    bullet_id = add_numbering_definition(doc, kind="bullet")
    for item in (
        "A clean machine can install domaininstall from npm and run di zuraai.xyz successfully.",
        "README, license, CI, package tarball, DNS record, and GitHub release all refer to the same tested version.",
        "Security limitations—especially first-use trust and normal npm install-script behavior—are stated plainly.",
        "All local smoke tests pass, the live E2E passes, and the published-package install test passes.",
        "The demo and documentation show only truthful, currently working examples.",
        "Rollback is known: deprecate the npm version, update/remove the TXT record, and publish a clear advisory.",
    ):
        add_list_item(doc, item, bullet_id)

    cadence_heading = add_heading(doc, "8. Operating cadence", 1)
    decimal_id = add_numbering_definition(doc, kind="decimal")
    for item in (
        "At the start of each work session, update only statuses whose evidence has changed.",
        "Before every merge, run npm test, package dry-run, TypeScript checks, and the relevant demo check.",
        "Before every release, run the live E2E and a clean registry installation—not only a local tarball install.",
        "After release, review metrics and trust objections weekly; move work by evidence, not by feature enthusiasm.",
    ):
        add_list_item(doc, item, decimal_id)

    add_callout(doc, "IMMEDIATE NEXT ACTION", "Reconcile main/default-branch state, add README + LICENSE + CI, then authenticate npm and publish domaininstall@0.0.1.", fill=GOLD_FILL, accent=GOLD_TEXT)

    # Core properties, then write.
    doc.core_properties.title = "domaininstall Project Execution Checklist"
    doc.core_properties.subject = "Release, security, and launch checklist"
    doc.core_properties.author = "domaininstall project"
    doc.core_properties.keywords = "domaininstall, di, DNS, npm, release checklist"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
